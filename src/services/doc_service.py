import json

from docx import Document
from docx.shared import Inches


class DocService:
    def __init__(self, customer_info: dict, company_info: dict, related_links: list):
        self.customer_info = customer_info.get("customer")
        self.company_info = company_info.get("company")
        self.related_links = related_links

    def create_doc(self):
        doc = Document()

        doc.add_heading("Customer Info", level=1)
        doc.add_paragraph("")
        self.__add_table(doc, self.customer_info)

        doc.add_paragraph("")

        doc.add_heading("Company Info", level=1)
        doc.add_paragraph("")
        self.__add_table(doc, self.company_info)

        doc.add_heading("Related links", level=1)
        doc.add_paragraph("")
        for link in self.related_links:
            doc.add_paragraph(link)

        doc.save(
            f"{self.customer_info.get('customer_name', 'unknown')} - {self.company_info.get('name', 'unknown')}.docx"
        )

    def __add_table(self, doc, info):
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        table.columns[0].width = Inches(2)
        table.columns[1].width = Inches(5)

        for key, value in info.items():
            row = table.add_row()
            row.cells[0].text = str(key)

            if isinstance(value, dict):
                for k, v in value.items():
                    row.cells[1].text += f"{k}: {v}\n"
            elif isinstance(value, list):
                for item in value:
                    row.cells[1].text += f"{item}\n"
            else:
                row.cells[1].text = str(value)

        return table
